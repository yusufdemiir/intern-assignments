from flask import Flask,  render_template, request, jsonify
import PyPDF2, xlrd, docx, os
from openpyxl import load_workbook
from datetime import datetime
import tempfile
from docx import Document
from PyPDF2 import PdfReader

app = Flask(__name__)

def get_metadata(file, filename):
    ext = os.path.splitext(filename)[1].upper()
    metadata = {}

    if ext == ".PDF":
        reader = PyPDF2.PdfReader(file)
        doc_info = reader.metadata
        creation_date = doc_info.get('/CreationDate')
        # Genellikle "D:YYYYMMDDHHMMSS" formatında gelir, başındaki "D:"'yi çıkar
        try:
            date_str = creation_date.strip().replace("D:", "")
            metadata["created"] = datetime.strptime(date_str[:14], "%Y%m%d%H%M%S")
        except Exception as e:
            metadata["created"] = f"Format hatası: {creation_date}"

        metadata["filename"] = filename
        metadata["type"] = ext
        metadata["author"] = doc_info.author
        metadata["producer"] = doc_info.producer

    elif ext == ".DOCX":
        doc = docx.Document(file)
        core_props = doc.core_properties
        metadata["filename"] = filename
        metadata["type"] = ext
        metadata["author"] = core_props.author
        metadata["created"] = str(core_props.created)
        metadata["producer"] = core_props.last_modified_by


    elif ext == ".TXT":
        with tempfile.NamedTemporaryFile(delete=False) as temp:
            temp.write(file.read())
            temp_path = temp.name

        stats = os.stat(temp_path)
        metadata["filename"] = filename
        metadata["type"] = ext
        metadata["author"] = " "
        metadata["created"] = datetime.fromtimestamp(stats.st_ctime)
        metadata["producer"] = " "
        os.remove(temp_path)


    elif ext == ".XLSX":
        with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as temp:
            temp.write(file.read())
            temp_path = temp.name
        wb = load_workbook(temp_path)
        props = wb.properties
        metadata["filename"] = filename
        metadata["type"] = ext
        metadata["author"] = props.creator
        metadata["created"] = datetime.fromtimestamp(os.path.getctime(temp_path))
        metadata["producer"] = props.lastModifiedBy
        os.remove(temp_path)
    else:
        metadata["error"] = "Desteklenmeyen dosya türü"

    return metadata

def extract_text(file_name):
    import os
    filename = file_name.filename 
    ext = os.path.splitext(filename)[1].lower()
    text = ""

    if ext == '.txt':
        text = file_name.read().decode('utf-8', errors='ignore')

    elif ext == '.docx':
        from docx import Document
        doc = Document(file_name)
        text = "\n".join([p.text for p in doc.paragraphs])

    elif ext == '.pdf':
        from PyPDF2 import PdfReader
        reader = PdfReader(file_name)
        text = "\n".join([page.extract_text() or "" for page in reader.pages])

    elif ext == '.xlsx':
        from openpyxl import load_workbook
        import tempfile
        import os

        # Geçici dosyaya kaydet
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp:
            file_name.save(tmp)
            tmp_path = tmp.name

        try:
            wb = load_workbook(tmp_path, read_only=True, data_only=True)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_text = [str(cell) if cell is not None else "" for cell in row]
                    text += " ".join(row_text) + "\n"
        finally:
            os.remove(tmp_path)

    return text

def search_in_text(text, query, search_type):
    results = []
    lines = text.splitlines()
    count = 0
    query_lower = query.lower()

    for idx, line in enumerate(lines, start=1):
        line_stripped = line.strip()
        line_lower = line_stripped.lower()

        if search_type == 'contains' and query_lower in line_lower:
            count += line_lower.count(query_lower)
            results.append((idx, line_stripped))

        elif search_type == 'startswith' and line_lower.startswith(query_lower):
            count += 1
            results.append((idx, line_stripped))

        elif search_type == 'endswith' and line_lower.endswith(query_lower):
            count += 1
            results.append((idx, line_stripped))

    return count, results

@app.route('/start', methods=["POST"])
def start():
    uploaded_files = request.files.getlist("files")
            
    results = []
    for file in uploaded_files:
        filename = file.filename
        metadata = get_metadata(file, filename)
        results.append(metadata)

    return render_template("index.html", data=results)

@app.route('/search', methods=["POST"])
def search():
    uploaded_files = request.files.getlist("files")
    query = request.form.get('query')
    search_type = request.form.get('search_type')  # 'contains', 'startswith', 'endswith'
    print(uploaded_files)
    print(query)
    print(search_type)
    results = []

    
    for file in uploaded_files:
        text = extract_text(file)
        count, matches = search_in_text(text, query, search_type)

        results.append({
            'filename': file.filename,
            'count': count,
            'matches': matches  # [(satır_no, satır_metni), ...]
        })
    return render_template('results.html', results=results, query=query, search_type=search_type)

@app.route('/')
def home():

    return render_template("index.html")

if __name__ == "__main__":
   app.run(debug=True)